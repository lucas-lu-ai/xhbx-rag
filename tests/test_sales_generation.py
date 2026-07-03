import base64
import asyncio
import inspect
import json
from io import BytesIO

import httpx
import pytest
from PIL import Image
from agentscope.message import DataBlock, TextBlock, ToolCallBlock
from agentscope.model import ChatResponse

from xhbx_rag.models import (
    CaseJourneyPart,
    CaseObjectionsPart,
    CaseSalesInsightsSource,
    CaseSalesScript,
    CaseSalesScriptDraft,
    CaseSalesStrategyDraft,
    CaseScriptsPart,
    CaseStrategiesPart,
    CustomerJourneyStepDraft,
    CustomerSignal,
    EvidenceRef,
    ObjectionEvidence,
    ObjectionHandlingDraft,
    SalesAction,
    SectionSalesEvidence,
    ScriptQuote,
    StrategyCandidate,
)
from xhbx_rag.sales_generation import (
    generate_case_sales_insights,
    generate_case_sales_insights_async,
)
from xhbx_rag.sales_generation import (
    SalesInsightAgentScopeAgent,
    SalesInsightGenerationError,
    VisionImageDescriptionAgent,
    _build_dashscope_chat_model,
    _build_structured_chat_model,
    _call_agent_scope_structured_output,
    _format_model_retry_error,
    _render_section_material,
    _RetryDiagnosticDashScopeChatModel,
    _RetryDiagnosticOpenAIChatModel,
)
from agentscope.credential import OpenAICredential
from xhbx_rag.observability import MemoryTraceSink
from xhbx_rag.source_loader import ParsedEmbeddedImage, ParsedSourceFile, SourceSection


class _FakeAgentScopeChatModel:
    def __init__(self, content: str) -> None:
        self.content = content
        self.calls: list[dict] = []

    async def __call__(self, messages, **kwargs):
        self.calls.append({"messages": messages, "kwargs": kwargs})
        return ChatResponse(
            content=[TextBlock(text=self.content)],
            is_last=True,
        )


class _FakeStructuredToolChatModel:
    def __init__(self, payload: dict) -> None:
        self.payload = payload
        self.calls: list[dict] = []

    async def __call__(self, messages, **kwargs):
        self.calls.append({"messages": messages, "kwargs": kwargs})
        return ChatResponse(
            content=[
                ToolCallBlock(
                    type="tool_call",
                    id="call-1",
                    name="generate_structured_output",
                    input=json.dumps(self.payload, ensure_ascii=False),
                )
            ],
            is_last=True,
        )


class _SequenceStructuredToolChatModel:
    def __init__(self, payloads: list[dict]) -> None:
        self.payloads = payloads
        self.calls: list[dict] = []

    async def __call__(self, messages, **kwargs):
        self.calls.append({"messages": messages, "kwargs": kwargs})
        payload = self.payloads[min(len(self.calls) - 1, len(self.payloads) - 1)]
        return ChatResponse(
            content=[
                ToolCallBlock(
                    type="tool_call",
                    id=f"call-{len(self.calls)}",
                    name="generate_structured_output",
                    input=json.dumps(payload, ensure_ascii=False),
                )
            ],
            is_last=True,
        )


class _RetryableDiagnosticChatModel(_RetryDiagnosticOpenAIChatModel):
    @classmethod
    def _get_retryable_exceptions(cls):
        return (RuntimeError,)

    async def _call_api(self, model_name, messages, **kwargs):
        if not hasattr(self, "attempt_count"):
            self.attempt_count = 0
        self.attempt_count += 1
        if self.attempt_count == 1:
            cause = ValueError("tcp reset by peer")
            raise RuntimeError("Connection error.") from cause
        return ChatResponse(content=[TextBlock(text="ok")], is_last=True)


class _StreamDisconnectThenSucceedChatModel(_RetryDiagnosticOpenAIChatModel):
    """第一次返回中途断连的流式响应，第二次返回完整流。"""

    async def _call_api(self, model_name, messages, **kwargs):
        if not hasattr(self, "attempt_count"):
            self.attempt_count = 0
        self.attempt_count += 1
        if self.attempt_count == 1:
            return self._broken_stream()
        return self._ok_stream()

    async def _broken_stream(self):
        yield ChatResponse(content=[TextBlock(text="部分输出")], is_last=False)
        raise httpx.RemoteProtocolError(
            "peer closed connection without sending complete message body"
        )

    async def _ok_stream(self):
        yield ChatResponse(content=[TextBlock(text="中间块")], is_last=False)
        yield ChatResponse(content=[TextBlock(text="完整回答")], is_last=True)


class _AlwaysDisconnectStreamChatModel(_RetryDiagnosticOpenAIChatModel):
    """每次调用都返回中途断连的流式响应。"""

    async def _call_api(self, model_name, messages, **kwargs):
        if not hasattr(self, "attempt_count"):
            self.attempt_count = 0
        self.attempt_count += 1
        return self._broken_stream()

    async def _broken_stream(self):
        yield ChatResponse(content=[TextBlock(text="部分输出")], is_last=False)
        raise httpx.RemoteProtocolError(
            "peer closed connection without sending complete message body"
        )


class _FakeHttpResponse:
    status_code = 502
    headers = {"x-request-id": "req_123", "content-type": "application/json"}
    text = '{"error":"upstream gateway closed"}'


class _FakeHttpStatusError(Exception):
    response = _FakeHttpResponse()


class _FakeSectionAgent:
    def extract(self, section):
        return SectionSalesEvidence(
            case_name=section.case_name,
            section_name=section.section_name,
            sales_actions=[
                SalesAction(
                    action="识别预算上限",
                    evidence="客户说每年不能超过80万",
                    source_refs=[
                        EvidenceRef(
                            section_name=section.section_name,
                            filename="第1节.track-0.txt",
                            quote="客户说每年不能超过80万",
                        )
                    ],
                )
            ],
        )


class _FakeCaseAgent:
    def __init__(self) -> None:
        self.received = None

    def extract(self, case_name, evidences):
        self.received = evidences
        assert evidences[0].sales_actions[0].source_refs[0].locator["line_start"] == 2
        return CaseSalesInsightsSource(
            case_name=case_name,
            case_summary="客户有预算上限，需要用预算释放方式处理。",
            scripts=[
                CaseSalesScript(
                    script_id="script_001",
                    stage="异议处理",
                    scenario="客户预算封顶",
                    customer_trigger="客户说每年不能超过80万",
                    goal="在预算不增加的前提下说明加保空间",
                    source_quote="客户说每年不能超过80万",
                    coach_wording="先确认预算红线，再检查已缴清保单释放出的预算。",
                    evidence_refs=[
                        EvidenceRef(
                            section_name="第1节",
                            filename="第1节.track-0.txt",
                            quote="客户说每年不能超过80万",
                        )
                    ],
                )
            ],
        )


def test_generate_case_sales_insights_writes_located_refs_and_playbook(tmp_path) -> None:
    case_dir = tmp_path / "案例A"
    section_dir = case_dir / "第1节"
    section_dir.mkdir(parents=True)
    (section_dir / "第1节.track-0.txt").write_text(
        "老师开场\n客户说每年不能超过80万\n销售回应：可以看缴费期满的保单\n",
        encoding="utf-8",
    )
    out_dir = tmp_path / "out"
    case_agent = _FakeCaseAgent()

    result = generate_case_sales_insights(
        case_dir=case_dir,
        output_dir=out_dir,
        section_agent=_FakeSectionAgent(),
        case_agent=case_agent,
    )

    assert result.status == "ok"
    data = json.loads(result.insights_path.read_text(encoding="utf-8"))
    ref = data["scripts"][0]["evidence_refs"][0]
    assert ref["source_type"] == "txt"
    assert ref["context"] == (
        "老师开场\n客户说每年不能超过80万\n销售回应：可以看缴费期满的保单"
    )
    assert ref["locator"]["line_start"] == 2
    assert ref["locator_confidence"] == "exact"
    evidence_data = json.loads(result.evidence_paths[0].read_text(encoding="utf-8"))
    source_ref = evidence_data["sales_actions"][0]["source_refs"][0]
    assert source_ref["context"] == ref["context"]
    assert result.playbook_path.exists()
    playbook = result.playbook_path.read_text(encoding="utf-8")
    assert "第1节.track-0.txt L2-L2" in playbook
    assert "客户说每年不能超过80万" in playbook


class _LenientFakeCaseAgent:
    def extract(self, case_name, evidences):
        return CaseSalesInsightsSource(
            case_name=case_name,
            case_summary="汇总",
            scripts=[
                CaseSalesScript(
                    script_id="script_001",
                    stage="异议处理",
                    scenario="客户预算封顶",
                    source_quote="客户说每年不能超过80万",
                    evidence_refs=[
                        EvidenceRef(
                            section_name="第1节",
                            filename="第1节.track-0.txt",
                            quote="客户说每年不能超过80万",
                        )
                    ],
                )
            ],
        )


class _ExplodingSectionAgent:
    def extract(self, section):
        raise AssertionError("复用模式下不应调用章节抽取")


def _write_reusable_section_evidence(out_dir, case_name, section_name) -> None:
    evidence = SectionSalesEvidence(
        case_name=case_name,
        section_name=section_name,
        sales_actions=[
            SalesAction(
                action="已有动作",
                evidence="已有证据",
                source_refs=[
                    EvidenceRef(
                        section_name=section_name,
                        filename="第1节.track-0.txt",
                        quote="客户说每年不能超过80万",
                    )
                ],
            )
        ],
    )
    case_out = out_dir / case_name
    case_out.mkdir(parents=True, exist_ok=True)
    (case_out / f"{section_name}.sales_evidence.json").write_text(
        json.dumps(evidence.model_dump(mode="json"), ensure_ascii=False),
        encoding="utf-8",
    )


def test_reuse_section_evidence_skips_section_agent(tmp_path) -> None:
    case_dir = tmp_path / "案例A"
    section_dir = case_dir / "第1节"
    section_dir.mkdir(parents=True)
    (section_dir / "第1节.track-0.txt").write_text(
        "老师开场\n客户说每年不能超过80万\n销售回应：可以看缴费期满的保单\n",
        encoding="utf-8",
    )
    out_dir = tmp_path / "out"
    _write_reusable_section_evidence(out_dir, "案例A", "第1节")
    trace = MemoryTraceSink()

    result = generate_case_sales_insights(
        case_dir=case_dir,
        output_dir=out_dir,
        section_agent=_ExplodingSectionAgent(),
        case_agent=_LenientFakeCaseAgent(),
        trace=trace,
        reuse_section_evidence=True,
    )

    assert result.status == "ok"
    steps = [event.step for event in trace.events]
    assert "generate.section_evidence_reused" in steps
    assert result.evidence_paths
    assert result.evidence_paths[0].name == "第1节.sales_evidence.json"


def test_reuse_section_evidence_falls_back_to_extraction_when_file_invalid(
    tmp_path,
) -> None:
    case_dir = tmp_path / "案例A"
    section_dir = case_dir / "第1节"
    section_dir.mkdir(parents=True)
    (section_dir / "第1节.track-0.txt").write_text(
        "老师开场\n客户说每年不能超过80万\n销售回应：可以看缴费期满的保单\n",
        encoding="utf-8",
    )
    out_dir = tmp_path / "out"
    case_out = out_dir / "案例A"
    case_out.mkdir(parents=True)
    (case_out / "第1节.sales_evidence.json").write_text(
        "{损坏的 JSON", encoding="utf-8"
    )
    trace = MemoryTraceSink()

    result = generate_case_sales_insights(
        case_dir=case_dir,
        output_dir=out_dir,
        section_agent=_FakeSectionAgent(),
        case_agent=_LenientFakeCaseAgent(),
        trace=trace,
        reuse_section_evidence=True,
    )

    assert result.status == "ok"
    steps = [event.step for event in trace.events]
    assert "generate.section_evidence_reused" not in steps
    assert "generate.section_evidence_written" in steps


class _FakePartCaseAgent:
    """支持分型调用的 fake 案例级 agent。"""

    def __init__(self, fail_parts: tuple[str, ...] = ()) -> None:
        self.calls: list[str] = []
        self.context_notes: dict[str, str] = {}
        self.fail_parts = set(fail_parts)

    def extract(self, case_name, evidences):
        raise AssertionError("分型模式下不应走单次大调用")

    async def extract_case_part_async(
        self, part, case_name, catalog_text, context_notes=""
    ):
        self.calls.append(part)
        self.context_notes[part] = context_notes
        if part in self.fail_parts:
            raise SalesInsightGenerationError(f"{part} 模拟失败")
        if part == "customer_journey":
            return CaseJourneyPart(
                case_summary="预算封顶客户的加保案例",
                customer_journey=[
                    CustomerJourneyStepDraft(
                        stage="异议处理",
                        customer_state="有预算顾虑",
                        sales_goal="确认预算红线",
                        key_actions=["确认红线"],
                        evidence_ids=["E001"],
                    )
                ],
            )
        if part == "strategies":
            return CaseStrategiesPart(
                strategies=[
                    CaseSalesStrategyDraft(
                        name="预算释放策略",
                        definition="用缴清保单释放预算空间",
                        confidence="high",
                        inferred=False,
                        evidence_ids=["E001"],
                    )
                ]
            )
        if part == "scripts":
            return CaseScriptsPart(
                scripts=[
                    CaseSalesScriptDraft(
                        script_id="script_001",
                        stage="异议处理",
                        scenario="客户预算封顶",
                        source_quote="客户说每年不能超过80万",
                        coach_wording="先确认红线，再看缴清保单释放的预算。",
                        strategy_names=["预算释放策略"],
                        evidence_ids=["E001", "E999"],
                    )
                ]
            )
        return CaseObjectionsPart(
            objection_handling=[
                ObjectionHandlingDraft(
                    objection="预算不能再增加",
                    diagnosis="客户以年度现金流设限",
                    recommended_response="用已缴清保单释放预算",
                    related_strategy_names=["预算释放策略"],
                    related_script_ids=["script_001"],
                    evidence_ids=["E001"],
                )
            ]
        )


def _make_case_material(tmp_path):
    case_dir = tmp_path / "案例A"
    section_dir = case_dir / "第1节"
    section_dir.mkdir(parents=True)
    (section_dir / "第1节.track-0.txt").write_text(
        "老师开场\n客户说每年不能超过80万\n销售回应：可以看缴费期满的保单\n",
        encoding="utf-8",
    )
    return case_dir


def test_case_parts_mode_assembles_full_insights_with_resolved_refs(tmp_path) -> None:
    case_dir = _make_case_material(tmp_path)
    out_dir = tmp_path / "out"
    agent = _FakePartCaseAgent()
    trace = MemoryTraceSink()

    result = generate_case_sales_insights(
        case_dir=case_dir,
        output_dir=out_dir,
        section_agent=_FakeSectionAgent(),
        case_agent=agent,
        trace=trace,
    )

    assert result.status == "ok"
    assert result.case_part_errors == ()
    assert agent.calls.count("customer_journey") == 1
    assert set(agent.calls) == {
        "customer_journey",
        "strategies",
        "scripts",
        "objection_handling",
    }
    assert "预算释放策略" in agent.context_notes["scripts"]
    assert "script_001" in agent.context_notes["objection_handling"]
    data = json.loads(result.insights_path.read_text(encoding="utf-8"))
    assert data["case_summary"] == "预算封顶客户的加保案例"
    script = data["scripts"][0]
    assert script["source_quote"] == "客户说每年不能超过80万"
    ref = script["evidence_refs"][0]
    assert ref["locator"]["line_start"] == 2
    assert ref["quote"] == "客户说每年不能超过80万"
    steps = [event.step for event in trace.events]
    assert steps.count("generate.case_part_extracted") == 4
    assert "generate.case_unknown_evidence_ids" in steps
    assert result.playbook_path.exists()


def test_case_parts_mode_partial_when_one_part_fails(tmp_path) -> None:
    case_dir = _make_case_material(tmp_path)
    out_dir = tmp_path / "out"
    agent = _FakePartCaseAgent(fail_parts=("strategies",))
    trace = MemoryTraceSink()

    result = generate_case_sales_insights(
        case_dir=case_dir,
        output_dir=out_dir,
        section_agent=_FakeSectionAgent(),
        case_agent=agent,
        trace=trace,
    )

    assert result.status == "partial"
    assert dict(result.case_part_errors).keys() == {"strategies"}
    data = json.loads(result.insights_path.read_text(encoding="utf-8"))
    assert data["strategies"] == []
    assert data["scripts"]
    assert data["objection_handling"]
    steps = [event.step for event in trace.events]
    assert "generate.case_part_failed" in steps


def test_case_parts_mode_failed_when_all_parts_fail(tmp_path) -> None:
    case_dir = _make_case_material(tmp_path)
    out_dir = tmp_path / "out"
    agent = _FakePartCaseAgent(
        fail_parts=("customer_journey", "strategies", "scripts", "objection_handling")
    )

    result = generate_case_sales_insights(
        case_dir=case_dir,
        output_dir=out_dir,
        section_agent=_FakeSectionAgent(),
        case_agent=agent,
    )

    assert result.status == "failed"
    assert result.insights_path is None
    assert len(dict(result.case_part_errors)) == 4


def test_case_parts_checkpoint_skips_completed_parts_on_rerun(tmp_path) -> None:
    case_dir = _make_case_material(tmp_path)
    out_dir = tmp_path / "out"

    first = generate_case_sales_insights(
        case_dir=case_dir,
        output_dir=out_dir,
        section_agent=_FakeSectionAgent(),
        case_agent=_FakePartCaseAgent(),
    )
    assert first.status == "ok"

    rerun_agent = _FakePartCaseAgent()
    trace = MemoryTraceSink()
    second = generate_case_sales_insights(
        case_dir=case_dir,
        output_dir=out_dir,
        section_agent=_FakeSectionAgent(),
        case_agent=rerun_agent,
        trace=trace,
    )

    assert second.status == "ok"
    assert rerun_agent.calls == []
    steps = [event.step for event in trace.events]
    assert steps.count("generate.case_part_reused") == 4


def test_case_parts_checkpoint_invalidated_when_evidence_changes(tmp_path) -> None:
    case_dir = _make_case_material(tmp_path)
    out_dir = tmp_path / "out"

    first = generate_case_sales_insights(
        case_dir=case_dir,
        output_dir=out_dir,
        section_agent=_FakeSectionAgent(),
        case_agent=_FakePartCaseAgent(),
    )
    assert first.status == "ok"

    class _OtherSectionAgent(_FakeSectionAgent):
        def extract(self, section):
            evidence = super().extract(section)
            return evidence.model_copy(
                update={
                    "sales_actions": [
                        evidence.sales_actions[0].model_copy(
                            update={"action": "改变后的动作"}
                        )
                    ]
                }
            )

    rerun_agent = _FakePartCaseAgent()
    second = generate_case_sales_insights(
        case_dir=case_dir,
        output_dir=out_dir,
        section_agent=_OtherSectionAgent(),
        case_agent=rerun_agent,
    )

    assert second.status == "ok"
    assert set(rerun_agent.calls) == {
        "customer_journey",
        "strategies",
        "scripts",
        "objection_handling",
    }


def test_case_parts_checkpoint_reruns_dependent_part_when_context_changes(
    tmp_path,
) -> None:
    case_dir = _make_case_material(tmp_path)
    out_dir = tmp_path / "out"

    first = generate_case_sales_insights(
        case_dir=case_dir,
        output_dir=out_dir,
        section_agent=_FakeSectionAgent(),
        case_agent=_FakePartCaseAgent(fail_parts=("strategies",)),
    )
    assert first.status == "partial"

    rerun_agent = _FakePartCaseAgent()
    second = generate_case_sales_insights(
        case_dir=case_dir,
        output_dir=out_dir,
        section_agent=_FakeSectionAgent(),
        case_agent=rerun_agent,
    )

    assert second.status == "ok"
    assert "strategies" in rerun_agent.calls
    assert "scripts" in rerun_agent.calls
    assert "objection_handling" in rerun_agent.calls
    assert "预算释放策略" in rerun_agent.context_notes["scripts"]
    data = json.loads(second.insights_path.read_text(encoding="utf-8"))
    assert data["scripts"][0]["strategy_names"] == ["预算释放策略"]


def test_case_parts_empty_result_is_not_frozen_into_checkpoint(tmp_path) -> None:
    case_dir = _make_case_material(tmp_path)
    out_dir = tmp_path / "out"

    class _EmptyStrategiesAgent(_FakePartCaseAgent):
        async def extract_case_part_async(
            self, part, case_name, catalog_text, context_notes=""
        ):
            if part == "strategies":
                self.calls.append(part)
                return CaseStrategiesPart(strategies=[])
            return await super().extract_case_part_async(
                part, case_name, catalog_text, context_notes
            )

    first = generate_case_sales_insights(
        case_dir=case_dir,
        output_dir=out_dir,
        section_agent=_FakeSectionAgent(),
        case_agent=_EmptyStrategiesAgent(),
    )
    assert first.status == "ok"
    assert not (out_dir / "案例A" / "case.insights_parts" / "strategies.json").exists()

    rerun_agent = _FakePartCaseAgent()
    second = generate_case_sales_insights(
        case_dir=case_dir,
        output_dir=out_dir,
        section_agent=_FakeSectionAgent(),
        case_agent=rerun_agent,
    )

    assert second.status == "ok"
    # strategies 重跑后产出了策略名单，scripts/objections 的输入随之变化，
    # part 指纹正确失效并重跑；只有 journey 与上游无依赖，仍复用 checkpoint。
    assert "customer_journey" not in rerun_agent.calls
    assert set(rerun_agent.calls) == {
        "strategies",
        "scripts",
        "objection_handling",
    }


def test_case_parts_checkpoint_write_failure_does_not_fail_the_part(
    tmp_path,
    monkeypatch,
) -> None:
    case_dir = _make_case_material(tmp_path)
    out_dir = tmp_path / "out"

    from xhbx_rag import sales_generation as sg

    def broken_write(path, fingerprint, part_result):
        raise OSError("磁盘已满")

    monkeypatch.setattr(sg, "_write_case_part_checkpoint", broken_write)

    result = generate_case_sales_insights(
        case_dir=case_dir,
        output_dir=out_dir,
        section_agent=_FakeSectionAgent(),
        case_agent=_FakePartCaseAgent(),
    )

    assert result.status == "ok"
    assert result.insights_path is not None


def test_reuse_section_evidence_rejects_mismatched_section_name(tmp_path) -> None:
    case_dir = _make_case_material(tmp_path)
    out_dir = tmp_path / "out"
    evidence = SectionSalesEvidence(
        case_name="案例A",
        section_name="完全不同的章节",
        sales_actions=[
            SalesAction(
                action="来自别的章节的动作",
                evidence="证据",
                source_refs=[EvidenceRef(filename="x.txt", quote="引用")],
            )
        ],
    )
    case_out = out_dir / "案例A"
    case_out.mkdir(parents=True)
    (case_out / "第1节.sales_evidence.json").write_text(
        json.dumps(evidence.model_dump(mode="json"), ensure_ascii=False),
        encoding="utf-8",
    )
    trace = MemoryTraceSink()

    result = generate_case_sales_insights(
        case_dir=case_dir,
        output_dir=out_dir,
        section_agent=_FakeSectionAgent(),
        case_agent=_LenientFakeCaseAgent(),
        trace=trace,
        reuse_section_evidence=True,
        case_call_mode="single",
    )

    assert result.status == "ok"
    steps = [event.step for event in trace.events]
    assert "generate.section_evidence_reused" not in steps
    assert "generate.section_evidence_written" in steps


def test_generate_rejects_unknown_case_call_mode(tmp_path) -> None:
    case_dir = _make_case_material(tmp_path)

    with pytest.raises(ValueError, match="case_call_mode"):
        generate_case_sales_insights(
            case_dir=case_dir,
            output_dir=tmp_path / "out",
            section_agent=_FakeSectionAgent(),
            case_agent=_FakePartCaseAgent(),
            case_call_mode="splitt",
        )


def test_case_call_mode_single_uses_legacy_case_agent_path(tmp_path) -> None:
    case_dir = _make_case_material(tmp_path)
    out_dir = tmp_path / "out"
    case_agent = _FakeCaseAgent()

    result = generate_case_sales_insights(
        case_dir=case_dir,
        output_dir=out_dir,
        section_agent=_FakeSectionAgent(),
        case_agent=case_agent,
        case_call_mode="single",
    )

    assert result.status == "ok"
    assert case_agent.received is not None


def test_sales_insight_agentscope_agent_uses_structured_tool_call() -> None:
    chat_model = _FakeStructuredToolChatModel(
        {"case_name": "案例A", "section_name": "第1节", "sales_actions": []}
    )
    agent = SalesInsightAgentScopeAgent(
        base_url="https://api.example.com/v1",
        api_key="secret",
        model="chat-model",
        chat_model=chat_model,
        retry_base_delay=0,
    )
    section = SourceSection(
        case_name="案例A",
        section_name="第1节",
        section_dir="案例A/第1节",
        sources=(
            ParsedSourceFile(
                source_id="txt:a.txt",
                source_type="txt",
                filename="a.txt",
                source_path="案例A/第1节/a.txt",
                text="客户关注预算",
            ),
        ),
    )

    agent.extract_section(section)

    call = chat_model.calls[0]
    assert "response_format" not in call["kwargs"]
    assert call["kwargs"]["tool_choice"].mode == "auto"
    assert call["kwargs"]["tools"][0]["function"]["name"] == "generate_structured_output"
    assert call["messages"][0].role == "system"
    assert call["messages"][1].role == "user"
    assert "客户关注预算" in call["messages"][1].get_text_content()
    assert "MUST" in call["messages"][1].get_text_content()


def test_retry_diagnostic_model_logs_exception_type_and_cause(monkeypatch) -> None:
    logs: list[str] = []

    def capture_warning(message, *args):
        logs.append(message % args)

    from xhbx_rag import sales_generation

    monkeypatch.setattr(sales_generation._agentscope_logger, "warning", capture_warning)
    model = _RetryableDiagnosticChatModel(
        credential=OpenAICredential(api_key="key", base_url="https://api.example.com/v1"),
        model="chat-model",
        max_retries=1,
        retry_delay=0,
    )

    response = asyncio.run(model([]))

    assert response.content[0].text == "ok"
    assert model.attempt_count == 2
    assert "诊断" in logs[0]
    assert "RuntimeError" in logs[0]
    assert "Connection error." in logs[0]
    assert "ValueError: tcp reset by peer" in logs[0]


def test_retry_log_message_signals_auto_recovery(monkeypatch) -> None:
    logs: list[str] = []

    def capture_warning(message, *args):
        logs.append(message % args)

    from xhbx_rag import sales_generation

    monkeypatch.setattr(sales_generation._agentscope_logger, "warning", capture_warning)
    model = _RetryableDiagnosticChatModel(
        credential=OpenAICredential(api_key="key", base_url="https://api.example.com/v1"),
        model="chat-model",
        max_retries=1,
        retry_delay=0,
    )

    asyncio.run(model([]))

    assert "自动重试" in logs[0]
    assert "任务不受影响" in logs[0]
    assert "1/2" in logs[0]
    assert "failed" not in logs[0]


def test_retry_exhausted_log_message_reports_final_failure(monkeypatch) -> None:
    logs: list[str] = []

    def capture_warning(message, *args):
        logs.append(message % args)

    from xhbx_rag import sales_generation

    monkeypatch.setattr(sales_generation._agentscope_logger, "warning", capture_warning)
    model = _AlwaysDisconnectStreamChatModel(
        credential=OpenAICredential(api_key="key", base_url="https://api.example.com/v1"),
        model="chat-model",
        stream=True,
        max_retries=1,
        retry_delay=0,
    )

    with pytest.raises(httpx.RemoteProtocolError):
        asyncio.run(model([]))

    assert "全部 2 次尝试均失败" in logs[-1]
    assert "诊断" in logs[-1]


def test_retry_mixin_drains_stream_and_retries_on_mid_stream_disconnect() -> None:
    model = _StreamDisconnectThenSucceedChatModel(
        credential=OpenAICredential(api_key="key", base_url="https://api.example.com/v1"),
        model="chat-model",
        stream=True,
        max_retries=1,
        retry_delay=0,
    )

    response = asyncio.run(model([]))

    assert model.attempt_count == 2
    assert not inspect.isasyncgen(response)
    assert response.content[0].text == "完整回答"


def test_retry_mixin_raises_after_stream_disconnect_exhausts_retries() -> None:
    model = _AlwaysDisconnectStreamChatModel(
        credential=OpenAICredential(api_key="key", base_url="https://api.example.com/v1"),
        model="chat-model",
        stream=True,
        max_retries=1,
        retry_delay=0,
    )

    with pytest.raises(httpx.RemoteProtocolError):
        asyncio.run(model([]))

    assert model.attempt_count == 2


def test_format_model_retry_error_includes_http_response_details() -> None:
    diagnostic = _format_model_retry_error(_FakeHttpStatusError("server failed"))

    assert "FakeHttpStatusError" in diagnostic
    assert "server failed" in diagnostic
    assert "status=502" in diagnostic
    assert "x-request-id=req_123" in diagnostic
    assert "body={\"error\":\"upstream gateway closed\"}" in diagnostic


def test_model_builders_use_retry_diagnostic_models() -> None:
    structured = _build_structured_chat_model(
        base_url="https://api.example.com/v1",
        api_key="key",
        model="chat-model",
        timeout=10,
        retry_attempts=2,
        retry_base_delay=0.25,
        enable_thinking=True,
        stream=True,
    )
    vision = _build_dashscope_chat_model(
        base_url="https://dashscope.example.com/compatible-mode/v1",
        api_key="key",
        model="vision-model",
        timeout=10,
        retry_attempts=2,
        retry_base_delay=0.25,
        enable_thinking=False,
        stream=True,
    )

    assert isinstance(structured, _RetryDiagnosticOpenAIChatModel)
    assert isinstance(vision, _RetryDiagnosticDashScopeChatModel)
    assert structured.max_retries == 1
    assert vision.max_retries == 1
    assert structured.stream is True
    assert vision.stream is True


def test_sales_insight_agent_prompt_includes_line_numbers_for_source_refs() -> None:
    chat_model = _FakeStructuredToolChatModel(
        {
            "case_name": "案例A",
            "section_name": "第1节",
            "sales_actions": [
                {
                    "action": "识别理赔价值",
                    "evidence": "客户一周内收到赔款",
                    "source_refs": [
                        {
                            "filename": "a.txt",
                            "quote": "他这90万的赔款在一周以内全部到账",
                            "locator": {"line_start": 2, "line_end": 4},
                        }
                    ],
                }
            ],
        }
    )
    agent = SalesInsightAgentScopeAgent(
        base_url="https://api.example.com/v1",
        api_key="secret",
        model="chat-model",
        chat_model=chat_model,
        retry_base_delay=0,
    )
    section = SourceSection(
        case_name="案例A",
        section_name="第1节",
        section_dir="案例A/第1节",
        sources=(
            ParsedSourceFile(
                source_id="txt:a.txt",
                source_type="txt",
                filename="a.txt",
                source_path="案例A/第1节/a.txt",
                text="开场白\n他这90万的赔款\n在一周以来\n全部到账\n",
            ),
        ),
    )

    evidence = agent.extract_section(section)

    user_text = chat_model.calls[0]["messages"][1].get_text_content()
    assert "L001: 开场白" in user_text
    assert "L002: 他这90万的赔款" in user_text
    assert "line_start/line_end" in user_text
    ref = evidence.sales_actions[0].source_refs[0]
    assert ref.locator["line_start"] == 2


def test_sales_insight_agent_can_compact_case_input() -> None:
    chat_model = _FakeStructuredToolChatModel(
        {"case_name": "案例A", "case_summary": "摘要"}
    )
    agent = SalesInsightAgentScopeAgent(
        base_url="https://api.example.com/v1",
        api_key="secret",
        model="chat-model",
        chat_model=chat_model,
        retry_base_delay=0,
        compact_case_input=True,
    )
    long_context = "上下文" * 300
    evidence = SectionSalesEvidence(
        case_name="案例A",
        section_name="第1节",
        customer_signals=[
            CustomerSignal(
                signal="客户担心企业资产传承",
                evidence="客户提到孩子接班和资产安全",
                source_refs=[
                    EvidenceRef(
                        section_name="第1节",
                        source_id="txt:a.txt",
                        source_type="txt",
                        filename="a.txt",
                        source_path="案例A/第1节/a.txt",
                        quote="孩子接班和资产安全",
                        context=long_context,
                        source_excerpt=long_context,
                        locator={"line_start": 7, "line_end": 8},
                        locator_confidence="exact",
                    )
                ],
            )
        ],
        sales_actions=[
            SalesAction(
                action="用财富观念唤醒传承需求",
                stage_hint="需求唤醒",
                evidence="销售引导客户区分企业资产和家庭资产",
            )
        ],
        script_quotes=[
            ScriptQuote(
                quote="先把企业资产和家庭资产分开看",
                speaker="销售",
                stage_hint="需求诊断",
            )
        ],
        objections=[
            ObjectionEvidence(
                objection="客户觉得保险收益不高",
                response_evidence="回应保险更看重确定性和传承安排",
            )
        ],
        strategy_candidates=[
            StrategyCandidate(
                name="财富观念唤醒",
                reason="先建立资产隔离与传承认知",
                confidence="high",
                inferred=False,
            )
        ],
    )

    agent.extract_case("案例A", [evidence])

    user_text = chat_model.calls[0]["messages"][1].get_text_content()
    assert "精简后的章节销售证据" in user_text
    assert "客户担心企业资产传承" in user_text
    assert "用财富观念唤醒传承需求" in user_text
    assert "先把企业资产和家庭资产分开看" in user_text
    assert "财富观念唤醒" in user_text
    assert '"line_start": 7' in user_text
    assert '"quote": "孩子接班和资产安全"' in user_text
    assert "context" not in user_text
    assert "source_excerpt" not in user_text
    assert long_context not in user_text


def test_generate_case_sales_insights_traces_final_case_model_input(tmp_path) -> None:
    case_dir = tmp_path / "案例A"
    section_dir = case_dir / "第1节"
    section_dir.mkdir(parents=True)
    (section_dir / "a.txt").write_text("客户关注预算", encoding="utf-8")
    chat_model = _FakeStructuredToolChatModel(
        {"case_name": "案例A", "case_summary": "预算相关摘要"}
    )
    case_agent = SalesInsightAgentScopeAgent(
        base_url="https://api.example.com/v1",
        api_key="secret",
        model="chat-model",
        chat_model=chat_model,
        retry_base_delay=0,
        compact_case_input=True,
    )
    trace = MemoryTraceSink()

    result = generate_case_sales_insights(
        case_dir=case_dir,
        output_dir=tmp_path / "out",
        section_agent=_FakeSectionAgent(),
        case_agent=case_agent,
        trace=trace,
        case_call_mode="single",
    )

    assert result.status == "ok"
    input_events = [
        event for event in trace.events if event.step == "generate.case_model_input"
    ]
    assert len(input_events) == 1
    payload = input_events[0].payload
    assert payload["case_name"] == "案例A"
    assert payload["evidence_count"] == 1
    assert payload["compact"] is True
    assert payload["model"] == "chat-model"
    assert payload["stream"] is False
    assert "案例级销售洞察专家" in payload["system_prompt"]
    assert "精简后的章节销售证据" in payload["user_content"]
    assert "识别预算上限" in payload["user_content"]
    assert payload["user_content_chars"] == len(payload["user_content"])
    model_user_text = chat_model.calls[0]["messages"][1].get_text_content()
    assert payload["user_content"] in model_user_text
    assert payload["structured_reminder"] in model_user_text


def test_agentscope_agent_case_part_call_uses_part_prompt_and_catalog() -> None:
    chat_model = _FakeStructuredToolChatModel(
        {
            "strategies": [
                {
                    "name": "预算释放策略",
                    "definition": "用缴清保单释放预算",
                    "evidence_ids": ["E001"],
                }
            ]
        }
    )
    agent = SalesInsightAgentScopeAgent(
        base_url="https://api.example.com/v1",
        api_key="secret",
        model="chat-model",
        chat_model=chat_model,
        retry_base_delay=0,
    )

    part = asyncio.run(
        agent.extract_case_part_async(
            "strategies",
            "案例A",
            "[E001] 识别预算上限｜解释：客户说每年不能超过80万",
            "本案例已确认策略名列表：预算释放策略",
        )
    )

    assert part.strategies[0].evidence_ids == ["E001"]
    call = chat_model.calls[0]
    system_text = call["messages"][0].get_text_content()
    user_text = call["messages"][1].get_text_content()
    assert "strategies" in system_text
    assert "跨章节去重合并" in system_text
    assert "[E001]" in user_text
    assert "本案例已确认策略名列表" in user_text


def test_sales_insight_agentscope_agent_unwraps_nested_structured_output() -> None:
    chat_model = _FakeStructuredToolChatModel(
        {
            "output": {
                "case_name": "案例A",
                "section_name": "第1节",
                "sales_actions": [
                    {
                        "action": "识别预算异议",
                        "evidence": "客户说每年不能超过80万",
                    }
                ],
            }
        }
    )
    agent = SalesInsightAgentScopeAgent(
        base_url="https://api.example.com/v1",
        api_key="secret",
        model="chat-model",
        chat_model=chat_model,
        retry_base_delay=0,
    )
    section = SourceSection(
        case_name="案例A",
        section_name="第1节",
        section_dir="案例A/第1节",
        sources=(
            ParsedSourceFile(
                source_id="txt:a.txt",
                source_type="txt",
                filename="a.txt",
                source_path="案例A/第1节/a.txt",
                text="客户说每年不能超过80万",
            ),
        ),
    )

    evidence = agent.extract_section(section)

    assert evidence.sales_actions[0].action == "识别预算异议"


def test_structured_output_retry_adds_validation_error_context() -> None:
    chat_model = _SequenceStructuredToolChatModel(
        [
            {"case_summary": "缺少必填 case_name"},
            {"case_name": "案例A", "case_summary": "摘要"},
        ]
    )

    data = _call_agent_scope_structured_output(
        chat_model,
        [
            TextBlock(text="system"),
            TextBlock(text="请生成案例洞察"),
        ],
        structured_model=CaseSalesInsightsSource,
    )

    assert data["case_name"] == "案例A"
    assert len(chat_model.calls) == 2
    second_user_text = chat_model.calls[1]["messages"][-1].get_text_content()
    assert "上一次结构化输出校验失败" in second_user_text
    assert "case_name" in second_user_text
    assert "缺少必填 case_name" in second_user_text


def test_vision_image_description_agent_sends_image_data_block() -> None:
    chat_model = _FakeAgentScopeChatModel("图片显示客户的保单整理表。")
    agent = VisionImageDescriptionAgent(
        base_url="https://api.example.com/v1",
        api_key="secret",
        model="qwen3.7-plus",
        chat_model=chat_model,
        retry_base_delay=0,
    )
    image = ParsedEmbeddedImage(
        image_id="docx:讲义.docx:image-1",
        filename="image1.png",
        source_path="案例A/第1节/讲义.docx::word/media/image1.png",
        media_type="image/png",
        data=b"png-bytes",
        locator={"container": "word/media/image1.png"},
    )
    source = ParsedSourceFile(
        source_id="docx:讲义.docx",
        source_type="docx",
        filename="讲义.docx",
        source_path="案例A/第1节/讲义.docx",
        text="讲义文字",
        images=(image,),
    )

    description = agent.describe(
        image,
        case_name="案例A",
        section_name="第1节",
        source=source,
    )

    assert description == "图片显示客户的保单整理表。"
    blocks = chat_model.calls[0]["messages"][1].get_content_blocks()
    data_blocks = [block for block in blocks if isinstance(block, DataBlock)]
    assert data_blocks
    assert data_blocks[0].source.media_type == "image/png"
    assert data_blocks[0].source.data == base64.b64encode(b"png-bytes").decode(
        "ascii"
    )


def test_generate_case_sales_insights_appends_image_descriptions(tmp_path) -> None:
    from docx import Document

    image_buffer = BytesIO()
    Image.new("RGB", (8, 8), "red").save(image_buffer, format="PNG")
    image_path = tmp_path / "chart.png"
    image_path.write_bytes(image_buffer.getvalue())
    case_dir = tmp_path / "案例A"
    section_dir = case_dir / "第1节"
    section_dir.mkdir(parents=True)
    docx_path = section_dir / "讲义.docx"
    document = Document()
    document.add_paragraph("这是一页保单整理讲义")
    document.add_picture(str(image_path))
    document.save(str(docx_path))

    class _VisionAgent:
        def describe(self, image, *, case_name, section_name, source):
            return "图片显示一张保单整理表，包含年缴保费和保障缺口。"

    class _SectionAgent:
        def __init__(self) -> None:
            self.source_text = ""

        def extract(self, section):
            self.source_text = section.sources[0].text
            return SectionSalesEvidence(
                case_name=section.case_name,
                section_name=section.section_name,
                sales_actions=[
                    SalesAction(
                        action="读取图片表格",
                        evidence="图片显示一张保单整理表",
                        source_refs=[
                            EvidenceRef(
                                filename="讲义.docx",
                                quote="图片显示一张保单整理表",
                            )
                        ],
                    )
                ],
            )

    class _CaseAgent:
        def extract(self, case_name, evidences):
            return CaseSalesInsightsSource(
                case_name=case_name,
                case_summary="摘要",
            )

    section_agent = _SectionAgent()

    result = generate_case_sales_insights(
        case_dir=case_dir,
        output_dir=tmp_path / "out",
        section_agent=section_agent,
        case_agent=_CaseAgent(),
        vision_agent=_VisionAgent(),
    )

    assert result.status == "ok"
    assert "图片补充信息" in section_agent.source_text
    assert "年缴保费和保障缺口" in section_agent.source_text


def test_generate_case_sales_insights_extracts_each_source_separately(tmp_path) -> None:
    case_dir = tmp_path / "案例A"
    section_dir = case_dir / "第1节"
    section_dir.mkdir(parents=True)
    (section_dir / "a.txt").write_text("客户关注预算", encoding="utf-8")
    (section_dir / "b.txt").write_text("销售建议预算释放", encoding="utf-8")

    class _SingleSourceAgent:
        def __init__(self) -> None:
            self.calls: list[list[str]] = []

        def extract(self, section):
            filenames = [source.filename for source in section.sources]
            self.calls.append(filenames)
            assert len(filenames) == 1
            return SectionSalesEvidence(
                case_name=section.case_name,
                section_name=section.section_name,
                sales_actions=[
                    SalesAction(
                        action=f"读取{filenames[0]}",
                        evidence=section.sources[0].text,
                        source_refs=[
                            EvidenceRef(
                                filename=filenames[0],
                                quote=section.sources[0].text,
                            )
                        ],
                    )
                ],
            )

    class _CaseAgent:
        def extract(self, case_name, evidences):
            assert len(evidences) == 1
            assert [item.action for item in evidences[0].sales_actions] == [
                "读取a.txt",
                "读取b.txt",
            ]
            return CaseSalesInsightsSource(
                case_name=case_name,
                case_summary="摘要",
            )

    section_agent = _SingleSourceAgent()

    result = generate_case_sales_insights(
        case_dir=case_dir,
        output_dir=tmp_path / "out",
        section_agent=section_agent,
        case_agent=_CaseAgent(),
    )

    assert result.status == "ok"
    assert section_agent.calls == [["a.txt"], ["b.txt"]]


def test_generate_case_sales_insights_async_reuses_one_loop_and_runs_sections_concurrently(
    tmp_path,
) -> None:
    case_dir = tmp_path / "案例A"
    for name in ["第1节", "第2节"]:
        section_dir = case_dir / name
        section_dir.mkdir(parents=True)
        (section_dir / f"{name}.txt").write_text(
            f"{name} 客户关注预算",
            encoding="utf-8",
        )

    class _AsyncSectionAgent:
        def __init__(self) -> None:
            self.loop_ids: list[int] = []
            self.active = 0
            self.max_active = 0

        async def extract_async(self, section):
            self.loop_ids.append(id(asyncio.get_running_loop()))
            self.active += 1
            self.max_active = max(self.max_active, self.active)
            await asyncio.sleep(0.01)
            self.active -= 1
            return SectionSalesEvidence(
                case_name=section.case_name,
                section_name=section.section_name,
                sales_actions=[
                    SalesAction(
                        action=f"读取{section.section_name}",
                        evidence=section.primary_text,
                        source_refs=[
                            EvidenceRef(
                                filename=section.sources[0].filename,
                                quote=section.primary_text.strip(),
                            )
                        ],
                    )
                ],
            )

    class _AsyncCaseAgent:
        def __init__(self) -> None:
            self.loop_ids: list[int] = []
            self.section_names: list[str] = []

        async def extract_async(self, case_name, evidences):
            self.loop_ids.append(id(asyncio.get_running_loop()))
            self.section_names = [evidence.section_name for evidence in evidences]
            return CaseSalesInsightsSource(case_name=case_name, case_summary="摘要")

    section_agent = _AsyncSectionAgent()
    case_agent = _AsyncCaseAgent()

    result = asyncio.run(
        generate_case_sales_insights_async(
            case_dir=case_dir,
            output_dir=tmp_path / "out",
            section_agent=section_agent,
            case_agent=case_agent,
            section_concurrency=2,
        )
    )

    assert result.status == "ok"
    assert section_agent.max_active == 2
    assert len(set(section_agent.loop_ids + case_agent.loop_ids)) == 1
    assert case_agent.section_names == ["第1节", "第2节"]


def test_generate_case_sales_insights_async_writes_failed_section_and_keeps_output(
    tmp_path,
) -> None:
    case_dir = tmp_path / "案例A"
    for name in ["第1节", "第2节"]:
        section_dir = case_dir / name
        section_dir.mkdir(parents=True)
        (section_dir / f"{name}.txt").write_text(
            f"{name} 客户关注预算",
            encoding="utf-8",
        )

    class _PartiallyFailingSectionAgent:
        async def extract_async(self, section):
            if section.section_name == "第2节":
                raise RuntimeError("模型返回内容无法修复")
            return SectionSalesEvidence(
                case_name=section.case_name,
                section_name=section.section_name,
                sales_actions=[
                    SalesAction(
                        action="识别预算",
                        evidence=section.primary_text,
                    )
                ],
            )

    class _CaseAgent:
        async def extract_async(self, case_name, evidences):
            assert [evidence.section_name for evidence in evidences] == ["第1节"]
            return CaseSalesInsightsSource(case_name=case_name, case_summary="摘要")

    result = asyncio.run(
        generate_case_sales_insights_async(
            case_dir=case_dir,
            output_dir=tmp_path / "out",
            section_agent=_PartiallyFailingSectionAgent(),
            case_agent=_CaseAgent(),
            section_concurrency=2,
        )
    )

    assert result.status == "ok"
    assert len(result.evidence_paths) == 1
    assert len(result.failure_paths) == 1
    failure = json.loads(result.failure_paths[0].read_text(encoding="utf-8"))
    assert failure["status"] == "failed"
    assert failure["section_name"] == "第2节"
    assert failure["error_type"] == "RuntimeError"
    assert "模型返回内容无法修复" in failure["error"]


def test_render_section_material_limits_prompt_size() -> None:
    section = SourceSection(
        case_name="案例A",
        section_name="第1节",
        section_dir="案例A/第1节",
        sources=(
            ParsedSourceFile(
                source_id="txt:a.txt",
                source_type="txt",
                filename="a.txt",
                source_path="案例A/第1节/a.txt",
                text="A" * 200,
            ),
        ),
    )

    material = _render_section_material(section, max_chars=120)

    assert len(material) <= 220
    assert "内容已截断" in material
